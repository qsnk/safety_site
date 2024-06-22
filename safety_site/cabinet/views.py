from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import logout
from django.core.cache import cache
from django.core.paginator import Paginator
from django.db.models import Count
from django.db.models.functions import TruncMonth, TruncDay, TruncHour
from django.utils.dateformat import DateFormat
from django.http.response import StreamingHttpResponse, HttpResponse, FileResponse
from cabinet.forms import AddCameraForm, AddPlaceForm, ShowPlaceForm, FilterJournalForm, AddReportForm, \
    FilterReportForm, FilterStatisticsForm, FilterCameraForm, FilterPlaceForm
from cabinet.models import Camera, Place, Violation, Report
from cabinet.camera import IpCamera
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from PIL import Image
from docx import Document
from docx.shared import Inches
from datetime import datetime
from urllib.parse import urljoin
import cv2, os, io, requests

# Create your views here.
@login_required(login_url="/login/")
def index(request):
    user = request.user
    username = user.username
    first_name = user.first_name
    last_name = user.last_name
    email = user.email
    date_joined = user.date_joined
    is_active = user.is_active
    user_info = [username, email, date_joined]
    context = {
        'user': user,
        'user_info': user_info,
        'username': username,
        'email': email,
        'first_name': first_name,
        'last_name': last_name,
        'date_joined': date_joined
    }
    return render(request, 'cabinet/index.html', context)

@login_required(login_url="/login/")
def log_out(request):
    logout(request)
    return redirect('/login/')

@login_required(login_url="/login/")
def add_cameras(request):
    cameras = Camera.objects.filter(user_id=request.user).order_by('-pk')

    if request.method == "POST":
        if 'add-camera-button' in request.POST:
            form = AddCameraForm(request.POST)
            if form.is_valid():
                name = form.cleaned_data["name"]
                url = form.cleaned_data["url"]
                user_id = request.user
                camera = Camera(name=name, url=url, user_id=user_id)
                camera.save()
            else:
                messages.error(request, "Invalid data!")

        if 'show-filter-button' in request.POST:
            form = FilterCameraForm(request.POST)
            if form.is_valid():
                date_start = form.cleaned_data["date_start"]
                date_end = form.cleaned_data["date_end"]
                month = form.cleaned_data["month"]
                time = form.cleaned_data["time"]

                if date_start:
                    cameras = cameras.filter(date_time__date=date_start).order_by('-pk')
                if date_end:
                    cameras = cameras.filter(date_time__date=date_end).order_by('-pk')
                if date_start and date_end:
                    cameras = cameras.filter(date_time__range=[date_start, date_end]).order_by('-pk')

                if month:
                    cameras = cameras.filter(date_time__year=month.split("-")[0],
                                                   date_time__month=month.split("-")[1]).order_by('-pk')
                if time:
                    cameras = cameras.filter(date_time__hour=time.hour, date_time__minute=time.minute).order_by(
                        '-pk')

                paginator = Paginator(cameras, 5)
                page_number = request.GET.get('page')
                page = paginator.get_page(page_number)
                context = {
                    'filter_form': form,
                    'cameras': cameras,
                    'page': page
                }
                return render(request, 'cabinet/add_cameras.html', context)

    add_form = AddCameraForm()
    filter_form = FilterCameraForm()
    paginator = Paginator(cameras, 5)
    page_number = request.GET.get('page')
    page = paginator.get_page(page_number)
    context = {
        'form': add_form,
        'filter_form': filter_form,
        'cameras': cameras,
        'page': page
    }
    return render(request, 'cabinet/add_cameras.html', context)

@login_required(login_url="/login/")
def add_places(request):
    places = Place.objects.filter(user_id=request.user).order_by('-pk')

    if request.method == "POST":
        if 'add-place-button' in request.POST:
            form = AddPlaceForm(request.POST, user_id=request.user)
            if form.is_valid():
                name = form.cleaned_data["name"]
                description = form.cleaned_data["description"]
                camera_id = form.cleaned_data["camera_id"]
                user_id = request.user
                place = Place(name=name, description=description, camera_id=camera_id, user_id=user_id)
                place.save()
            else:
                messages.error(request, "Invalid data!")

        if 'show-filter-button' in request.POST:
            form = FilterPlaceForm(request.POST)
            if form.is_valid():
                date_start = form.cleaned_data["date_start"]
                date_end = form.cleaned_data["date_end"]
                month = form.cleaned_data["month"]
                time = form.cleaned_data["time"]
                if date_start:
                    places = places.filter(date_time__date=date_start).order_by('-pk')
                if date_end:
                    places = places.filter(date_time__date=date_end).order_by('-pk')
                if date_start and date_end:
                    places = places.filter(date_time__range=[date_start, date_end]).order_by('-pk')

                if month:
                    places = places.filter(date_time__year=month.split("-")[0],
                                             date_time__month=month.split("-")[1]).order_by('-pk')
                if time:
                    places = places.filter(date_time__hour=time.hour, date_time__minute=time.minute).order_by(
                        '-pk')

                paginator = Paginator(places, 5)
                page_number = request.GET.get('page')
                page = paginator.get_page(page_number)
                context = {
                    'filter_form': form,
                    'places': places,
                    'page': page
                }
                return render(request, 'cabinet/add_place.html', context)

    add_form = AddPlaceForm(user_id=request.user)
    filter_form = FilterPlaceForm()
    paginator = Paginator(places, 5)
    page_number = request.GET.get('page')
    page = paginator.get_page(page_number)
    context = {
        'form': add_form,
        'filter_form': filter_form,
        'places': places,
        'page': page
    }
    return render(request, 'cabinet/add_place.html', context)

@login_required(login_url="/login/")
def watch_site(request):
    sites_cache = cache.get('sites')
    print(sites_cache)
    if request.method == "POST":
        form = ShowPlaceForm(request.POST, user_id=request.user)
        if 'stop-button' in request.POST:
            cache.set('sites', None, 60 * 60)
            cv2.VideoCapture().release()
            cv2.VideoWriter().release()
            form = ShowPlaceForm(user_id=request.user)
            context = {'form': form}
            return render(request, 'cabinet/watch_site.html', context)

        if 'show-button' in request.POST:
            if form.is_valid():
                choices = form.cleaned_data["places"]
                if choices:
                    sites = {place.pk: place.name for place in choices}
                    cache.set('sites', sites, 60 * 60 * 24)
                    print(sites)
                else:
                    sites = None
                    cache.set('sites', sites, 60 * 60)
                context = {'form': form, 'sites': sites}
                return render(request, 'cabinet/watch_site.html', context)

    form = ShowPlaceForm(user_id=request.user)
    places = Place.objects.filter(user_id=request.user)
    context = {'places': places, 'sites': sites_cache, 'form': form}
    return render(request, 'cabinet/watch_site.html', context)

def stream_video(request, id):
    sites_cached = cache.get('sites')
    if sites_cached is not None:
        chosen_camera = Place.objects.get(pk=id).camera_id
        url = chosen_camera.url
        camera = IpCamera(url)
        return StreamingHttpResponse(generate_frames(request, camera), content_type='multipart/x-mixed-replace; boundary=frame')
    else:
        sites = None
        cache.set('sites', sites, 60 * 60)
        form = ShowPlaceForm(user_id=request.user)
        places = Place.objects.filter(user_id=request.user)
        context = {'sites': sites, 'form': form, 'places': places}
        return render(request, 'cabinet/watch_site.html', context)

@login_required(login_url="/login/")
def journal(request):
    violations = Violation.objects.filter(user_id=request.user).order_by('-pk')

    if request.method == "POST":
        form = FilterJournalForm(request.POST, user_id=request.user)
        if form.is_valid():
            date_start = form.cleaned_data["date_start"]
            date_end = form.cleaned_data["date_end"]
            month = form.cleaned_data["month"]
            time = form.cleaned_data["time"]
            violation_classes = form.cleaned_data["violations"]

            if date_start:
                violations = violations.filter(date_time__date=date_start).order_by('-pk')
            if date_end:
                violations = violations.filter(date_time__date=date_end).order_by('-pk')
            if date_start and date_end:
                violations = violations.filter(date_time__range=[date_start, date_end]).order_by('-pk')

            if month:
                violations = violations.filter(date_time__year=month.split("-")[0], date_time__month=month.split("-")[1]).order_by('-pk')
            if time:
                violations = violations.filter(date_time__hour=time.hour, date_time__minute=time.minute).order_by('-pk')

            if violation_classes:
                violations = violations.filter(violation_class_ru__in=violation_classes).order_by('-pk')

            paginator = Paginator(violations, 20)
            page_number = request.GET.get('page')
            page = paginator.get_page(page_number)
            context = {
                'form': form,
                'date_start': date_start,
                'date_end': date_end,
                'time': time,
                'violation_classes': violation_classes,
                'page': page
            }
            return render(request, 'cabinet/journal.html', context)

        else:
            print("Invalid data!")
            messages.error(request, "Invalid data!")

    date_start = request.GET.get('date_start')
    date_end = request.GET.get('date_end')
    time = request.GET.get('time')
    violation_classes = request.GET.get('violation_classes')
    form = FilterJournalForm(user_id=request.user)
    paginator = Paginator(violations, 20)
    page_number = request.GET.get('page')
    page = paginator.get_page(page_number)
    context = {
        'violations': violations,
        'form': form,
        'page': page,
        'date_start': date_start,
        'date_end': date_end,
        'time': time,
        'violation_classes': violation_classes,
    }
    return render(request, 'cabinet/journal.html', context)

@login_required(login_url="/login/")
def reports(request):
    reports = Report.objects.filter(user_id=request.user).order_by('-pk')

    if request.method == "POST":
        if 'add-report-button' in request.POST:
            form = AddReportForm(request.POST, user_id=request.user)
            if form.is_valid():
                name = form.cleaned_data["name"]
                date_time = datetime.now()
                violation_id = form.cleaned_data["violation"]
                user_id = request.user
                report_filename = f'{name}-{date_time.day}-{date_time.month}-{date_time.year}-{date_time.hour}-{date_time.minute}.pdf'
                report_file_path = os.path.join('reports', report_filename)
                file = report_file_path
                report = Report(name=name, date_time=date_time.strftime('%d.%m.%y %H:%M'), violation_id=violation_id, file=file, user_id=user_id)
                report.save()
                create_pdf_report(
                    request=request,
                    title=name,
                    date_time=date_time.strftime('%d.%m.%y %H:%M'),
                    violation_object=violation_id,
                    image_path=violation_id.photo.url,
                    output_path=os.path.join('media', 'reports',
                        f'{name}-{datetime.now().day}-{datetime.now().month}-{datetime.now().year}-{datetime.now().hour}-{datetime.now().minute}.pdf')
                )

        if 'show-filter-button' in request.POST:
            filter_form = FilterReportForm(request.POST, user_id=request.user)
            if filter_form.is_valid():
                date_start = filter_form.cleaned_data["date_start"]
                date_end = filter_form.cleaned_data["date_end"]
                time = filter_form.cleaned_data["time"]
                violation_classes = filter_form.cleaned_data["violations"]

                if date_start:
                    reports = Report.objects.filter(date_time__date=date_start).order_by('-pk')
                if date_end:
                    reports = Report.objects.filter(date_time__date=date_end).order_by('-pk')
                if date_start and date_end:
                    reports = Report.objects.filter(date_time__range=[date_start, date_end]).order_by('-pk')

                if time:
                    reports = Report.objects.filter(date_time__hour=time.hour, date_time__minute=time.minute).order_by('-pk')

                if violation_classes:
                    violations = Violation.objects.filter(violation_class_ru__in=violation_classes)
                    reports = Report.objects.filter(violation_id__in=violations).order_by('-pk')

                report_form = AddReportForm(user_id=request.user)
                paginator = Paginator(reports, 20)
                page_number = request.GET.get('page')
                page = paginator.get_page(page_number)
                context = {
                    'report_form': report_form,
                    'filter_form': filter_form,
                    'date_start': date_start,
                    'date_end': date_end,
                    'time': time,
                    'violation_classes': violation_classes,
                    'page': page
                }
                return render(request, 'cabinet/reports.html', context)

            else:
                print("Invalid data!")
                messages.error(request, "Invalid data!")

    report_form = AddReportForm(user_id=request.user)
    filter_form = FilterReportForm(user_id=request.user)
    paginator = Paginator(reports, 10)
    page_number = request.GET.get('page')
    page = paginator.get_page(page_number)
    context = {
        'reports': reports,
        'page': page,
        'report_form': report_form,
        'filter_form': filter_form
    }
    return render(request, 'cabinet/reports.html', context)

@login_required(login_url="/login/")
def statistics(request):
    if request.method == "POST":
        form = FilterStatisticsForm(request.POST)
        if form.is_valid():
            day = form.cleaned_data["day"]
            month_year = form.cleaned_data["month_year"]
            data = None
            violations_count = None
            if day:
                violations = Violation.objects.filter(date_time__day=day.day).order_by('-pk')
                hourly_violations = violations.filter(date_time__date=day) \
                    .annotate(hour=TruncHour('date_time')) \
                    .values('hour') \
                    .annotate(count=Count('id')) \
                    .order_by('hour')

                data = {
                    'labels': [entry['hour'].strftime('%H:00') for entry in hourly_violations],
                    'counts': [entry['count'] for entry in hourly_violations]
                }
                violations_count = len(violations)
            if month_year:
                year = month_year.split("-")[0]
                month = month_year.split("-")[1]
                violations = Violation.objects.filter(user_id=request.user).filter(date_time__year=year, date_time__month=month).order_by('-pk')
                daily_violations = violations.filter(date_time__year=year, date_time__month=month) \
                    .annotate(day=TruncDay('date_time')) \
                    .values('day') \
                    .annotate(count=Count('id')) \
                    .order_by('day')

                data = {
                    'labels': [entry['day'].strftime('%d %B %Y') for entry in daily_violations],
                    'counts': [entry['count'] for entry in daily_violations]
                }

                violations_count = len(violations)

            context = {
                'form': form,
                'day': day,
                'month_year': month_year,
                'violations_count': violations_count,
                'violations': data,
                'data': data,
                'zip_data': zip(data['labels'], data['counts']),

            }
            return render(request, 'cabinet/statistics.html', context)

    form = FilterStatisticsForm()
    violations = Violation.objects.filter(user_id=request.user).order_by('-pk')
    violations_count = len(violations)
    months = violations.annotate(month=TruncMonth('date_time')).values('month').annotate(count=Count('id')).order_by('month')
    data = [{'month': DateFormat(entry['month']).format('F Y'), 'count': entry['count']} for entry in months]
    context = {
        'form': form,
        'all_violations': violations,
        'violations_count': violations_count,
        'data': data,
        'months': [item["month"] for item in data]
    }
    return render(request, 'cabinet/statistics.html', context)

def generate_frames(request, camera):
    while camera.capture.isOpened():
        frame = camera.get_frame(request)
        yield (b'--frame\r\n'
                b'Content-Type: image/jpeg\r\n\r\n' + frame + b'\r\n\r\n')

def create_pdf_report(request, title, date_time, violation_object, image_path, output_path):
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter

    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase import pdfmetrics
    pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
    c.setFont('DejaVuSans', 24)
    c.drawString(1 * inch, height - 1 * inch, title)
    c.setTitle(title)
    c.setFont('DejaVuSans', 12)
    c.drawString(1 * inch, height - 1.5 * inch, f"Дата и время формирования: {date_time}")
    c.drawString(1 * inch, height - 2 * inch, f"Объект нарушения: {violation_object}")
    c.drawString(1 * inch, height - 2.5 * inch, "Изображение:")
    full_image_url = urljoin(f'{request.scheme}://{request.get_host()}', image_path)

    try:
        response = requests.get(full_image_url)
        image = ImageReader(io.BytesIO(response.content))
        image_width, image_height = image.getSize()
        aspect = image_height / float(image_width)
        display_width = 4 * inch
        display_height = display_width * aspect
        if display_height > (height - 3 * inch):
            display_height = height - 3 * inch
            display_width = display_height / aspect
        c.drawImage(image, 1 * inch, height - 3 * inch - display_height, display_width, display_height)
    except Exception as e:
        c.drawString(1 * inch, height - 3 * inch, "Изображение не может быть загружено.")
        c.drawString(1 * inch, height - 3.5 * inch, str(e))

    c.showPage()
    c.save()

def create_word_report(request, title, date_time, violation_object, image_url):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Дата и время формирования: {date_time}")
    doc.add_paragraph(f"Объект нарушения: {violation_object}")
    doc.add_paragraph("Изображение:")
    full_image_url = urljoin(f'{request.scheme}://{request.get_host()}', image_url)

    try:
        response = requests.get(full_image_url)
        image = Image.open(io.BytesIO(response.content))
        image_stream = io.BytesIO()
        image.save(image_stream, format='PNG')
        image_stream.seek(0)
        doc.add_picture(image_stream, width=Inches(4))
    except Exception as e:
        doc.add_paragraph("Изображение не может быть загружено.")
        doc.add_paragraph(str(e))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@login_required(login_url="/login/")
def download_word_report(request, report_id):
    report = get_object_or_404(Report, pk=report_id)
    buffer = create_word_report(
        request=request,
        title=report.name,
        date_time=report.date_time.strftime('%d.%m.%y %H:%M'),
        violation_object=report.violation_id,
        image_url=report.violation_id.photo.url
    )

    response = FileResponse(buffer, as_attachment=True, filename=f'{report.name}.docx')
    return response